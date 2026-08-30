from docx import Document
from pathlib import Path

ref = Path(r"C:\Users\arach\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-design-report\assets\reference.docx")
out = Path(r"C:\Users\arach\Documents\Projects\SIH26147\SIH26147_design_report.docx")
d = Document(ref)

repl = {
2: "SIH26147 Project Design Report",
7: "SIH26147 is a Python-based RF signal-forensics system for blind signal recovery, modulation identification, synchronization, data reconstruction, and scientific verification.",
8: "The project is architected as a seven-phase pipeline, supported by CLI tools, a desktop/web UI, sample IQ datasets, reproducible reports, and a large regression test suite.",
10: "15,000+ lines across approximately 153 Python files",
11: "77 test files covering DSP, modulation, recovery, FEC, verification, orchestration, and reporting",
12: "Current environment cannot run tests: pytest is absent from the project virtual environment and system Python is policy-blocked",
14: "The repository follows a clear measurement-to-recovery model: physical signal measurements are separated from modulation hypotheses, recovery attempts, and independent verification.",
15: "The scope is ambitious and research-oriented, but operational readiness depends on dependency reproducibility, repository hygiene, and stronger failure visibility.",
17: "The strongest design decision is the explicit governing invariant: Phase 2 measures, Phase 3 hypothesizes, Phase 4 recovers, Phase 5 corrects, and Phase 6 verifies.",
19: "The codebase is modular, with dedicated packages for IO, DSP, modulation, recovery, data recovery, verification, orchestration, reporting, and UI.",
21: "The worktree currently includes generated caches, egg-info, uploads, and modified report artifacts, indicating that development outputs are not fully isolated from source control.",
23: "Key takeaway. This is a strong research prototype with a coherent end-to-end architecture; the next priority should be making execution and validation reproducible.",
26: "Without a verified test run, claims about production readiness remain provisional. GUI fallbacks and broad exception handling may also conceal missing dependencies or runtime defects.",
28: "Prioritize environment reproducibility, test execution, source-control hygiene, and an end-to-end smoke test before adding major analytical features.",
29: "Make setup deterministic. Add all GUI and reporting dependencies as explicit optional extras and document the supported installation path.",
30: "Stabilize validation. Install and run pytest in the bundled or project environment, record the actual test count, and add a representative example-IQ smoke test.",
31: "Improve maintainability. Expand .gitignore coverage, remove generated artifacts from tracked workflows, and replace broad exception catches with targeted errors and diagnostics.",
33: "The project has a credible foundation for scientific signal analysis and recovery. Its architecture is more mature than its current delivery hygiene.",
34: "A focused hardening pass should convert the existing prototype into a reliable, repeatable analysis tool.",
37: "Evidence reviewed: README.md, pyproject.toml, docs/architecture.md, repository structure, git status, and available test inventory.",
38: "Execution note: pytest validation was attempted but blocked by the local environment; no test-pass claim is made in this report.",
40: "Project source. SIH26147 repository. Local workspace inspection, 29 August 2026.",
41: "Project documentation. Architecture and phase guides. Local workspace, 29 August 2026.",
42: "Project test suite. tests/ directory and example IQ datasets. Local workspace, 29 August 2026.",
43: "Environment diagnostic. Project virtual environment and system Python availability. Local workspace, 29 August 2026.",
45: "Report note. This assessment reflects the repository state inspected on 29 August 2026; generated artifacts and existing user changes were preserved.",
}
for i, text in repl.items():
    d.paragraphs[i].text = text

t = d.tables[1]
rows = [
    ("Architecture", "Seven-phase pipeline from ingestion through verification and reporting", "Clear separation of measurement, hypothesis, recovery, and validation"),
    ("Coverage", "Broad module and test inventory across DSP, modulation, recovery, and UI", "Good foundation for regression testing and extension"),
    ("Readiness", "Tests could not run in the current environment; generated files pollute the worktree", "Reproducibility and delivery hygiene are the immediate risks"),
]
for row, vals in zip(t.rows[1:], rows):
    for cell, value in zip(row.cells, vals):
        cell.text = value

d.save(out)
print(out)
